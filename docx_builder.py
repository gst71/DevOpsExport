"""
Word-Generator für Detailkonzept
--------------------------------
Lädt die bossinfo.ch Detailkonzept-Vorlage NUR als Style-Träger
(Schriften, Heading-Styles, kommentarintern, Header/Footer, Logo),
leert den Inhalt komplett und baut frische Kapitel aus DevOps Work Items.
"""
from __future__ import annotations

import io
import os
import re
from datetime import date
from html.parser import HTMLParser
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL

from devops_client import WorkItem, AzureDevOpsClient


# -------- HTML -> Word -------- #

class _HtmlToDocx(HTMLParser):
    """
    Minimaler HTML-Parser, der DevOps-Descriptions (kommt als HTML) sinnvoll
    in einen Word-Bereich überführt. Unterstützt:
      - <p>, <br>, <div>          -> Absätze
      - <b>/<strong>, <i>/<em>, <u> -> Inline-Formatierung
      - <ul>/<ol>/<li>            -> Listen
      - <h1>..<h6>                -> Inline fett (kein neuer Heading-Style,
                                     damit unsere Kapitel-Struktur stabil bleibt)
      - <img src="...">           -> Wenn src in image_map: Bild einfügen
      - <a href="...">            -> Link-Text fett
      - <table>/<tr>/<td>         -> einfache Tabelle
    """

    def __init__(self, doc, image_map: dict, devops_client, ref_paragraph=None):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.image_map = image_map
        self.devops_client = devops_client
        self.ref_paragraph = ref_paragraph  # falls Inhalt vor einem Marker eingefügt werden soll
        self.current_p = None
        self.bold = False
        self.italic = False
        self.underline = False
        self.in_list = []   # Stack: 'ul' / 'ol'
        self.list_index = []
        self.in_table = False
        self.table = None
        self.row_cells = None
        self.cell_idx = 0
        # Anzahl aufeinanderfolgender <br> im aktuellen Absatz - damit
        # <br><br><br>... auf maximal 1 Zeilenumbruch reduziert wird.
        self._br_streak = 0
        # Wurde im aktuellen Absatz schon echter Text/ein Bild eingefuegt?
        self._p_has_real_content = False

    def _ensure_paragraph(self):
        if self.current_p is None:
            self.current_p = self.doc.add_paragraph()
            # Den standardmaessigen Word-Abstand nach jedem Absatz entfernen,
            # damit der Output 1:1 wie im Work Item aussieht (DevOps verwendet
            # tighte HTML-Zeilenabstaende; eine Leerzeile = eine Leerzeile).
            pf = self.current_p.paragraph_format
            pf.space_after = Pt(0)
            pf.space_before = Pt(0)
            # Linienabstand explizit auf Single setzen (Normal-Style hat 1.08x)
            try:
                from docx.enum.text import WD_LINE_SPACING
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            except Exception:
                pass
            # Absatzmarken-Schriftgroesse auf 10pt setzen, damit leere
            # Absaetze exakt eine Body-Text-Zeile hoch sind (statt 11pt
            # aus dem Normal-Style).
            self._set_para_mark_size(self.current_p, 10)
        return self.current_p

    @staticmethod
    def _set_para_mark_size(paragraph, size_pt: int) -> None:
        """Setzt die Schriftgroesse der Absatzmarke (massgeblich fuer leere Absaetze)."""
        from docx.oxml import OxmlElement
        pPr = paragraph._p.get_or_add_pPr()
        # Existierendes rPr in pPr finden/erzeugen
        rPr = pPr.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            pPr.append(rPr)
        # sz/szCs setzen (Werte sind halb-Punkte)
        for tag, val in (("w:sz", size_pt * 2), ("w:szCs", size_pt * 2)):
            existing = rPr.find(qn(tag))
            if existing is not None:
                rPr.remove(existing)
            el = OxmlElement(tag)
            el.set(qn("w:val"), str(val))
            rPr.append(el)
        # Schriftart auf Arial fuer konsistente Hoehe
        existing = rPr.find(qn("w:rFonts"))
        if existing is None:
            f = OxmlElement("w:rFonts")
            f.set(qn("w:ascii"), "Arial")
            f.set(qn("w:hAnsi"), "Arial")
            rPr.append(f)

    def _close_paragraph(self):
        self.current_p = None
        self._br_streak = 0
        self._p_has_real_content = False

    def _add_run(self, text: str):
        if not text:
            return
        # Whitespace-only Daten ueberspringen, wenn noch kein Absatz existiert
        # (verhindert leere Spacer-Absaetze aus <p>&nbsp;</p> etc.)
        if self.current_p is None and not text.strip():
            return
        # Sobald echter Text reinkommt, ist der Absatz "gefuellt" -
        # nachfolgende <br> zaehlen wieder ab 0.
        if text.strip():
            self._br_streak = 0
            self._p_has_real_content = True
        p = self._ensure_paragraph()
        run = p.add_run(text)
        run.bold = self.bold
        run.italic = self.italic
        run.underline = self.underline
        # Standard-Schrift Arial 10
        run.font.name = "Arial"
        run.font.size = Pt(10)

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        a = dict(attrs)

        if tag in ("p", "div"):
            # Jeder <p>/<div>-Block erhaelt einen Absatz. Empty-Tags ergeben
            # eine Leerzeile (1:1 wie im Work Item). Die spaetere Cleanup-Pass
            # reduziert mehrere aufeinander folgende Leerzeilen auf eine.
            self._close_paragraph()
            self._ensure_paragraph()
        elif tag == "br":
            # Fuehrende <br> ignorieren
            if self.current_p is None:
                return
            # Wenn der Absatz noch keinen Inhalt hat, fuegen wir KEINEN
            # zusaetzlichen Umbruch ein. Der Absatz selbst ist ja schon
            # eine Leerzeile - <br> obendrauf wuerde die Hoehe verdoppeln.
            if not self._p_has_real_content:
                return
            # Mehrere <br> hintereinander auf maximal einen Umbruch reduzieren
            if self._br_streak >= 1:
                return
            self._br_streak += 1
            self.current_p.add_run().add_break()
        elif tag in ("b", "strong"):
            self.bold = True
        elif tag in ("i", "em"):
            self.italic = True
        elif tag == "u":
            self.underline = True
        elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._close_paragraph()
            self.bold = True
        elif tag == "ul":
            self.in_list.append("ul")
            self.list_index.append(0)
        elif tag == "ol":
            self.in_list.append("ol")
            self.list_index.append(0)
        elif tag == "li":
            self._close_paragraph()
            try:
                style = "List Bullet" if self.in_list and self.in_list[-1] == "ul" else "List Number"
                self.current_p = self.doc.add_paragraph(style=style)
            except KeyError:
                self.current_p = self.doc.add_paragraph()
                self.current_p.add_run("• " if (self.in_list and self.in_list[-1] == "ul") else "")
        elif tag == "table":
            self.in_table = True
            self._close_paragraph()
            # Vorerst zweispaltige Standard-Tabelle (wird dynamisch erweitert)
            self.table = self.doc.add_table(rows=0, cols=0)
            self.table.style = "Table Grid"
        elif tag == "tr":
            if self.table is not None:
                # Spaltenzahl im ersten <tr> ermitteln (siehe handle_endtag tr)
                self.row_cells = []
                self.cell_idx = 0
        elif tag in ("td", "th"):
            if self.in_table:
                self.row_cells.append([])  # Sammelt Text-Fragmente
        elif tag == "a":
            self.bold = True
        elif tag == "img":
            src = a.get("src", "")
            self._insert_image(src, a)

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in ("p", "div"):
            self._close_paragraph()
        elif tag in ("b", "strong"):
            self.bold = False
        elif tag in ("i", "em"):
            self.italic = False
        elif tag == "u":
            self.underline = False
        elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self.bold = False
            self._close_paragraph()
        elif tag in ("ul", "ol"):
            if self.in_list:
                self.in_list.pop()
                self.list_index.pop()
            self._close_paragraph()
        elif tag == "li":
            self._close_paragraph()
        elif tag == "table":
            self.in_table = False
            self.table = None
        elif tag == "tr":
            if self.table is not None and self.row_cells is not None:
                # Reihe in Tabelle anlegen
                ncols = len(self.row_cells)
                if ncols == 0:
                    return
                # Spalten ggf. erweitern
                while len(self.table.columns) < ncols:
                    self.table.add_column(Inches(1.2))
                row = self.table.add_row()
                for i, frag in enumerate(self.row_cells):
                    if i < len(row.cells):
                        row.cells[i].text = "".join(frag).strip()
                self.row_cells = None
        elif tag in ("td", "th"):
            pass
        elif tag == "a":
            self.bold = False

    def handle_data(self, data):
        if self.in_table and self.row_cells is not None and len(self.row_cells) > 0:
            self.row_cells[-1].append(data)
        else:
            self._add_run(data)

    # --------- Bilder ---------

    def _insert_image(self, src: str, attrs: dict):
        """
        Bettet ein Bild ein. Robustes Verhalten:
          - Akzeptiert data:-URLs, absolute und relative URLs
          - Skaliert grosse Bilder automatisch auf max. 16cm Breite
          - Bei Fehler wird ein roter Platzhalter-Text eingefuegt statt
            stilles Verschlucken, damit man im Output sieht, was fehlt.
        """
        if not src:
            self._image_placeholder("[Bild ohne src-Attribut]")
            return

        img_bytes: bytes | None = None

        # 1) Schon vorab geladene Bilder
        if src in self.image_map:
            img_bytes = self.image_map[src]
        # 2) Bilder via DevOps-Client laden (oder andere Quellen)
        elif self.devops_client is not None and hasattr(self.devops_client, "fetch_image"):
            img_bytes = self.devops_client.fetch_image(src)

        if not img_bytes:
            self._image_placeholder(f"[Bild konnte nicht geladen werden: {src[:80]}]")
            return

        # Schliesse aktuellen Absatz, fuege Bild in neuem Absatz ein
        self._close_paragraph()
        p = self.doc.add_paragraph()
        # Kein extra Abstand vor/nach Bildabsatz
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        run = p.add_run()

        # Bildbreite bestimmen: maximal 16cm, sonst original
        width = self._compute_image_width(img_bytes, attrs)

        ok = False
        for try_attempt in range(2):
            try:
                if width:
                    run.add_picture(io.BytesIO(img_bytes), width=width)
                else:
                    run.add_picture(io.BytesIO(img_bytes))
                ok = True
                break
            except Exception:
                # Fallback: Falls Format nicht erkannt wird, mit PIL konvertieren
                if try_attempt == 0:
                    converted = self._try_convert_to_png(img_bytes)
                    if converted:
                        img_bytes = converted
                        continue
                break

        if not ok:
            # Bild-Run entfernen und Platzhalter zeigen
            self._image_placeholder(f"[Bild im Format nicht unterstuetzt: {src[:80]}]")

    def _compute_image_width(self, img_bytes: bytes, attrs: dict):
        """Maximale Breite 16cm, schmaler wenn das Bild kleiner ist."""
        max_width_cm = 16.0
        # Versuche, aus dem Bild die Originaldimension zu ermitteln
        try:
            from PIL import Image
            img = Image.open(io.BytesIO(img_bytes))
            w_px = img.width
            # Annahme 96 DPI
            w_cm = w_px / 96.0 * 2.54
            if w_cm < max_width_cm:
                return Cm(w_cm)
            return Cm(max_width_cm)
        except Exception:
            pass
        # Falls PIL nicht verfuegbar oder Bild nicht lesbar, defaults
        return Cm(max_width_cm)

    def _try_convert_to_png(self, img_bytes: bytes) -> bytes | None:
        """Versucht, das Bild via Pillow zu PNG zu konvertieren (fuer TIFF, WebP, etc.)."""
        try:
            from PIL import Image
            img = Image.open(io.BytesIO(img_bytes))
            if img.mode not in ("RGB", "RGBA"):
                img = img.convert("RGBA")
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            return buf.getvalue()
        except Exception:
            return None

    def _image_placeholder(self, message: str) -> None:
        """Sichtbarer roter Platzhalter bei Bild-Fehler (damit nichts heimlich fehlt)."""
        self._close_paragraph()
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        run = p.add_run(message)
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


# -------- Hauptklasse -------- #

class DetailkonzeptBuilder:
    """
    Baut das vollständige Detailkonzept-Word aus einem Epic-Hierarchie-Baum.
    Verwendet die Vorlage NUR als Style-Träger.
    """

    def __init__(self, template_path: str, devops_client: AzureDevOpsClient | None = None,
                 allowed_types: set[str] | None = None,
                 show_work_item_id: bool = True,
                 show_estimates_section: bool = True):
        self.template_path = template_path
        self.devops_client = devops_client
        # Wenn None oder leer: keine Filterung (alle Typen werden gedruckt).
        # Wenn gesetzt: nur Work Items dieser Typen werden ins Word aufgenommen.
        # Epic ist immer enthalten, da es die Wurzel ist.
        self.allowed_types: set[str] | None = (
            set(allowed_types) if allowed_types else None
        )
        # Steuert, ob hinter jedem Titel die Work-Item-ID als (#NNNN)-Hyperlink
        # erscheint. Wenn False, werden Titel "clean" ohne ID gedruckt.
        self.show_work_item_id = show_work_item_id
        # Steuert, ob das Kapitel "Aufwandschätzung Entwicklungen" mit der
        # Estimate-Tabelle vor dem "Freigabe"-Feature eingefuegt wird.
        self.show_estimates_section = show_estimates_section

    def _type_allowed(self, wi_type: str) -> bool:
        if not self.allowed_types:
            return True
        return wi_type in self.allowed_types

    # ----- Public API -----

    def build(self, epic_tree: WorkItem, customer_name: str, output_path: str) -> str:
        doc = Document(self.template_path)
        self._clear_body(doc)

        self._add_title_page(doc, epic_tree, customer_name)
        self._add_page_break(doc)
        self._add_toc_placeholder(doc)
        self._add_page_break(doc)

        self._add_epic_section(doc, epic_tree)

        # [Kunde] global ersetzen
        if customer_name:
            self._replace_placeholder(doc, "[Kunde]", customer_name)

        # Mehrfache leere Absaetze auf maximal einen reduzieren
        self._cleanup_empty_paragraphs(doc)

        doc.save(output_path)
        return output_path

    def _cleanup_empty_paragraphs(self, doc: Document) -> None:
        """
        Bereinigt ueberfluessige Leerabsaetze:
          1. Aufeinanderfolgende Leere -> auf maximal einen reduzieren
          2. Leerabsaetze direkt vor einer Ueberschrift entfernen
             (Heading-Style hat eigenen Abstand davor)
          3. Leerabsaetze direkt nach einer Ueberschrift entfernen
             (Heading-Style hat eigenen Abstand danach)
          4. Trailing Leerzeilen am Dokumentende entfernen
        """
        body = doc.element.body
        children = [el for el in list(body) if el.tag == qn("w:p")]

        def is_heading(el) -> bool:
            ppr = el.find(qn("w:pPr"))
            if ppr is None:
                return False
            pstyle = ppr.find(qn("w:pStyle"))
            if pstyle is None:
                return False
            val = pstyle.get(qn("w:val"), "")
            return val.startswith("Heading") or val in ("Titel", "Title")

        def is_empty(el) -> bool:
            text = "".join((t.text or "") for t in el.iter(qn("w:t")))
            text = text.replace("\xa0", " ").strip()
            has_image = el.find(".//" + qn("w:drawing")) is not None
            return (not text) and (not has_image)

        # Mehrere Durchgaenge, bis nichts mehr entfernt werden kann
        while True:
            to_remove = []
            for i, el in enumerate(children):
                if not is_empty(el):
                    continue
                prev_el = children[i - 1] if i > 0 else None
                next_el = children[i + 1] if i < len(children) - 1 else None

                # (1) zwei aufeinanderfolgende leere
                if prev_el is not None and is_empty(prev_el):
                    to_remove.append(el)
                    continue
                # (2) leer direkt vor einer Heading
                if next_el is not None and is_heading(next_el):
                    to_remove.append(el)
                    continue
                # (3) leer direkt nach einer Heading
                if prev_el is not None and is_heading(prev_el):
                    to_remove.append(el)
                    continue
                # (4) trailing am Ende (kein nachfolgender Inhalt mehr)
                if next_el is None:
                    to_remove.append(el)
                    continue

            if not to_remove:
                break
            for el in to_remove:
                if el.getparent() is body:
                    body.remove(el)
            children = [el for el in list(body) if el.tag == qn("w:p")]

    # ----- Body leeren -----

    def _clear_body(self, doc: Document) -> None:
        """Entfernt alle Inhalts-Elemente (Paragraphs/Tables) im Body,
        lässt aber sectPr (Header/Footer-Verknüpfung, Page-Setup) intakt."""
        body = doc.element.body
        for child in list(body):
            if child.tag == qn("w:sectPr"):
                continue
            body.remove(child)
        # Mindestens einen leeren Paragraph behalten, sonst macht Word Stress
        doc.add_paragraph()

    # ----- Title Page -----

    def _add_title_page(self, doc: Document, epic: WorkItem, customer: str) -> None:
        """
        Titelseite ohne leere Spacer-Paragraphs aufbauen.
        Stattdessen werden Vor-/Nach-Abstaende auf den eigentlichen Absaetzen
        gesetzt - so kann die globale Cleanup-Routine die Titelseite nicht
        versehentlich kollabieren lassen.
        """
        def add_centered(text: str, size_pt: int, bold: bool = False,
                         space_before_pt: int = 0, space_after_pt: int = 0,
                         color: RGBColor | None = None) -> None:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            if space_before_pt:
                pf.space_before = Pt(space_before_pt)
            if space_after_pt:
                pf.space_after = Pt(space_after_pt)
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(size_pt)
            run.bold = bold
            if color is not None:
                run.font.color.rgb = color

        # Titel
        add_centered("Detailkonzept", size_pt=28, bold=True,
                     space_before_pt=130, space_after_pt=24)
        # Projekt / Epic
        add_centered(epic.title or f"Epic {epic.id}", size_pt=18, bold=True,
                     space_before_pt=12, space_after_pt=24)
        # Kunde
        add_centered(
            customer or "[Kunde]", size_pt=14, bold=False,
            space_before_pt=12, space_after_pt=0,
            color=None if customer else RGBColor(0xC0, 0x00, 0x00),
        )
        # Datum / Version - weiter unten auf der Seite
        add_centered(date.today().strftime("%d. %B %Y") + "  –  Version 1.0",
                     size_pt=11, bold=False,
                     space_before_pt=220, space_after_pt=0)
        # Footer-Zeile bossinfo.ch AG
        add_centered("bossinfo.ch AG", size_pt=11, bold=True,
                     space_before_pt=24, space_after_pt=0)

    # ----- TOC -----

    def _add_toc_placeholder(self, doc: Document) -> None:
        p = doc.add_paragraph()
        run = p.add_run("Inhalt")
        run.font.name = "Arial"
        run.font.size = Pt(20)
        run.bold = True

        # TOC-Feld: Word ergänzt es beim ersten Öffnen (Felder aktualisieren)
        p = doc.add_paragraph()
        run = p.add_run()
        fldChar1 = _make_element("w:fldChar", {"w:fldCharType": "begin"})
        instrText = _make_element("w:instrText", {"xml:space": "preserve"},
                                   text=r' TOC \o "1-3" \h \z \u ')
        fldChar2 = _make_element("w:fldChar", {"w:fldCharType": "separate"})
        fldChar3 = _make_element("w:t", {}, text="Inhaltsverzeichnis wird beim Öffnen aktualisiert (F9).")
        fldChar4 = _make_element("w:fldChar", {"w:fldCharType": "end"})
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        run._r.append(fldChar4)

    # ----- Inhalt aus Work Items -----

    def _add_epic_section(self, doc: Document, epic: WorkItem) -> None:
        # Optional: Epic-Beschreibung als Einleitung
        if epic.description_html.strip():
            self._add_heading(doc, "Einleitung", level=1)
            self._add_html(doc, epic.description_html)
        if epic.acceptance_html.strip():
            self._add_heading(doc, "Zielsetzung", level=1)
            self._add_html(doc, epic.acceptance_html)

        # Features = H1
        features = [c for c in epic.children if c.work_item_type == "Feature"] or epic.children
        estimates_emitted = False
        for feature in features:
            # Vor "Freigabe"-Feature die Aufwandschaetzungs-Tabelle einfuegen
            if (self.show_estimates_section
                    and not estimates_emitted
                    and "freigabe" in (feature.title or "").lower()):
                self._add_estimates_section(doc, epic)
                estimates_emitted = True

            # Type-Filter: ist dieser Typ aktiv?
            if not self._type_allowed(feature.work_item_type):
                continue

            self._add_heading(doc, feature.title, level=1,
                              work_item_id=feature.id, work_item_url=feature.url)
            if feature.description_html.strip():
                self._add_html(doc, feature.description_html)
            if feature.acceptance_html.strip():
                self._add_kommentar_intern(doc, "Acceptance Criteria")
                self._add_html(doc, feature.acceptance_html)

            # Stories = H2 (auch PBI/Issue/Bug auf gleicher Ebene möglich)
            stories = [c for c in feature.children
                       if c.work_item_type in ("User Story", "Product Backlog Item",
                                               "Issue", "Bug", "Requirement")]
            if not stories:
                stories = feature.children
            for story in stories:
                if not self._type_allowed(story.work_item_type):
                    continue
                self._add_heading(doc, story.title, level=2,
                                  work_item_id=story.id, work_item_url=story.url)
                if story.description_html.strip():
                    self._add_html(doc, story.description_html)
                if story.acceptance_html.strip():
                    self._add_kommentar_intern(doc, "Acceptance Criteria")
                    self._add_html(doc, story.acceptance_html)

                # Tasks = H3
                for task in story.children:
                    if not self._type_allowed(task.work_item_type):
                        continue
                    self._add_heading(doc, task.title, level=3,
                                      work_item_id=task.id, work_item_url=task.url)
                    if task.description_html.strip():
                        self._add_html(doc, task.description_html)
                    if task.acceptance_html.strip():
                        self._add_kommentar_intern(doc, "Acceptance Criteria")
                        self._add_html(doc, task.acceptance_html)

    # ----- Helpers -----

    def _add_heading(self, doc: Document, text: str, level: int,
                     work_item_id: int | None = None,
                     work_item_url: str | None = None) -> None:
        """
        Erzeugt eine Ueberschrift im exakten bossinfo-Stil:
        Verwendet die eingebauten Heading-1/2/3-Styles der Vorlage (schwarz,
        Arial Black/Arial, fett, mit automatischer 1 / 1.1 / 1.1.1-Nummerierung).
        Falls work_item_id+url uebergeben werden, wird " (#NNNN)" als
        anklickbarer Hyperlink hinter den Titel gehaengt.
        """
        # Heading-Level begrenzen auf die in der Vorlage definierten Stile
        level = max(1, min(level, 9))
        p = doc.add_paragraph()
        try:
            p.style = doc.styles["Heading %d" % level]
        except KeyError:
            # Notfall-Fallback: nur dann hart kodieren, wenn der Style fehlt
            run = p.add_run(text or "(ohne Titel)")
            run.font.name = "Arial"
            run.bold = True
            run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
            return

        # Titel-Text: KEIN font/size/bold setzen - das macht der Style
        p.add_run(text or "(ohne Titel)")

        # Work-Item-ID als klickbarer Hyperlink anhaengen (nur wenn aktiviert)
        if self.show_work_item_id and work_item_id and work_item_url:
            # Heading-Schriftgroessen passend zum Level (10pt fuer alle Levels)
            link_size = 10
            # Klammern als normale Runs (nicht bold) damit sie nicht im
            # Heading-Style-Fett-Look erscheinen
            opener = p.add_run(" (")
            opener.font.name = "Arial"
            opener.font.size = Pt(link_size)
            opener.bold = False
            _add_hyperlink_to_paragraph(
                p, f"#{work_item_id}", work_item_url,
                color="0563C1", underline=True, font="Arial", size_pt=link_size,
            )
            closer = p.add_run(")")
            closer.font.name = "Arial"
            closer.font.size = Pt(link_size)
            closer.bold = False

    # ----- Aufwandschaetzung -----

    # Bekannte Feldnamen fuer Original Estimate (DEV) und Original Estimate PL.
    # Die genauen Custom-Feldnamen koennen je nach Prozess-Template unterschiedlich sein,
    # daher pruefen wir mehrere Kandidaten und matchen auch via Patternsuche.
    _DEV_ESTIMATE_FIELDS = (
        "Microsoft.VSTS.Scheduling.OriginalEstimate",
        "Custom.OriginalEstimate",
        "Custom.OriginalEstimateDev",
    )
    _PL_ESTIMATE_FIELDS = (
        "Microsoft.VSTS.Scheduling.OriginalEstimatePL",
        "Custom.OriginalEstimatePL",
        "Custom.OriginalEstimateProjectLead",
    )

    def _get_estimate(self, fields: dict, candidates: tuple, fallback_pattern: str = "") -> float:
        """
        Versucht aus den DevOps-Feldern den Aufwand zu extrahieren.
        Erst werden die bekannten Feldnamen geprueft, dann ein Pattern-Match
        ueber alle vorhandenen Felder.
        """
        # 1) Bekannte Feldnamen
        for name in candidates:
            v = fields.get(name)
            if v is not None:
                try:
                    return float(v)
                except (TypeError, ValueError):
                    continue
        # 2) Pattern-Suche: irgendein Feld dessen Name den Pattern matcht
        if fallback_pattern:
            pl = fallback_pattern.lower()
            for key, val in fields.items():
                kl = key.lower()
                if pl in kl and "originalestimate" in kl:
                    try:
                        return float(val)
                    except (TypeError, ValueError):
                        continue
        return 0.0

    def _collect_estimates(self, root: WorkItem) -> list[tuple[WorkItem, float, float]]:
        """
        Sammelt alle Work Items, bei denen mindestens eines der beiden
        Estimate-Felder einen Wert > 0 hat. Reihenfolge folgt der Hierarchie
        (also Backlog-Reihenfolge wie im Rest des Dokuments).
        """
        result: list[tuple[WorkItem, float, float]] = []

        def walk(node: WorkItem):
            dev = self._get_estimate(node.raw_fields, self._DEV_ESTIMATE_FIELDS)
            pl = self._get_estimate(node.raw_fields, self._PL_ESTIMATE_FIELDS,
                                    fallback_pattern="pl")
            # Mindestens einer der beiden Werte muss > 0 sein
            # und der Typ muss im Filter erlaubt sein.
            if ((dev and dev > 0) or (pl and pl > 0)) and self._type_allowed(node.work_item_type):
                result.append((node, dev, pl))
            for child in node.children:
                walk(child)

        walk(root)
        return result

    def _add_estimates_section(self, doc: Document, epic: WorkItem) -> None:
        """Eigenes Kapitel mit Aufwandschaetzungs-Tabelle."""
        self._add_heading(doc, "Aufwandschätzung Entwicklungen", level=1)

        estimates = self._collect_estimates(epic)
        if not estimates:
            p = doc.add_paragraph(
                "(Keine Work Items mit gesetztem Original Estimate gefunden.)"
            )
            for r in p.runs:
                r.font.name = "Arial"; r.font.size = Pt(10); r.italic = True
            return

        # Tabelle: 4 Spalten - ID/Typ (mit Hyperlink) | Beschreibung | Aufwand DEV | Aufwand PL
        table = doc.add_table(rows=1, cols=4)
        try:
            table.style = doc.styles["Table Grid"]
        except KeyError:
            pass

        # Spaltenbreiten (Summe ~16cm = A4 Schreibbereich)
        widths = [Cm(3.2), Cm(7.3), Cm(2.7), Cm(2.7)]
        for col, w in zip(table.columns, widths):
            for cell in col.cells:
                cell.width = w

        def style_cell(cell, text: str, bold: bool = False, align_right: bool = False):
            cell.text = ""
            p = cell.paragraphs[0]
            if align_right:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.bold = bold
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        def hyperlink_cell(cell, text: str, url: str):
            """Erste Spalte: ID/Typ als anklickbarer Hyperlink."""
            cell.text = ""
            p = cell.paragraphs[0]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if url:
                _add_hyperlink_to_paragraph(p, text, url,
                                            color="0563C1", underline=True,
                                            font="Arial", size_pt=10)
            else:
                run = p.add_run(text)
                run.font.name = "Arial"
                run.font.size = Pt(10)

        # Header
        hdr = table.rows[0].cells
        style_cell(hdr[0], "Work Item", bold=True)
        style_cell(hdr[1], "Beschreibung", bold=True)
        style_cell(hdr[2], "Aufwand DEV", bold=True, align_right=True)
        style_cell(hdr[3], "Aufwand PL", bold=True, align_right=True)
        for c in hdr:
            _shade_paragraph_cell(c, "DDDDDD")

        # Kompakte Typ-Anzeige
        type_short = {
            "User Story": "Story",
            "Product Backlog Item": "PBI",
        }

        total_dev = 0.0
        total_pl = 0.0
        for wi, dev, pl in estimates:
            row = table.add_row().cells
            type_label = type_short.get(wi.work_item_type, wi.work_item_type or "?")
            id_label = f"#{wi.id} {type_label}".strip()
            hyperlink_cell(row[0], id_label, wi.url or "")
            style_cell(row[1], wi.title or f"#{wi.id}")
            style_cell(row[2], f"{dev:.1f} h" if dev else "", align_right=True)
            style_cell(row[3], f"{pl:.1f} h" if pl else "", align_right=True)
            total_dev += dev
            total_pl += pl

        # Summenzeile
        sum_row = table.add_row().cells
        style_cell(sum_row[0], "", bold=True)
        style_cell(sum_row[1], "Total", bold=True)
        style_cell(sum_row[2], f"{total_dev:.1f} h", bold=True, align_right=True)
        style_cell(sum_row[3], f"{total_pl:.1f} h", bold=True, align_right=True)
        for c in sum_row:
            _shade_paragraph_cell(c, "F2F2F2")

    def _add_kommentar_intern(self, doc: Document, label: str) -> None:
        """
        Box-artige Hervorhebung im 'kommentar intern'-Stil
        (Courier New, weisser Text auf dunkelblauem Hintergrund).
        """
        p = doc.add_paragraph()
        try:
            p.style = doc.styles["kommentar intern"]
        except KeyError:
            pass
        run = p.add_run(label)
        run.font.name = "Courier New"
        run.font.size = Pt(10.5)
        run.bold = True
        # Falls Style nicht vorhanden: manuelle Farben
        try:
            doc.styles["kommentar intern"]
        except KeyError:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_paragraph(p, "1F4E79")

    def _add_html(self, doc: Document, html: str) -> None:
        if not html:
            return
        # DevOps speichert Bilder manchmal als Markdown (`![alt](url)`) statt
        # als <img>. Vorab in echte img-Tags konvertieren.
        html = _markdown_images_to_html(html)
        # Manchmal sind auch Links als Markdown drin: [text](url) -> <a>
        html = _markdown_links_to_html(html)
        parser = _HtmlToDocx(doc, image_map={}, devops_client=self.devops_client)
        parser.feed(html)
        parser.close()

    def _add_page_break(self, doc: Document) -> None:
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

    def _replace_placeholder(self, doc: Document, placeholder: str, value: str) -> None:
        # In allen Paragraphen (auch Header/Footer) ersetzen
        def replace_in_paragraphs(paragraphs):
            for p in paragraphs:
                if placeholder in p.text:
                    inline = p.runs
                    full_text = "".join(r.text for r in inline)
                    if placeholder in full_text:
                        new_text = full_text.replace(placeholder, value)
                        # Alles in den ersten Run packen, restliche leeren
                        if inline:
                            inline[0].text = new_text
                            for r in inline[1:]:
                                r.text = ""
        replace_in_paragraphs(doc.paragraphs)
        for section in doc.sections:
            replace_in_paragraphs(section.header.paragraphs)
            replace_in_paragraphs(section.footer.paragraphs)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)


# -------- Markdown-Vorverarbeitung -------- #

_MD_IMG_RE = re.compile(r"!\[([^\]]*)\]\(([^)\s]+(?:\([^)]*\)[^)\s]*)*)\)")
_MD_LINK_RE = re.compile(r"(?<!!)\[([^\]]+)\]\(([^)\s]+(?:\([^)]*\)[^)\s]*)*)\)")


def _markdown_images_to_html(text: str) -> str:
    """Konvertiert Markdown-Bilder ![alt](url) in echte <img>-Tags."""
    if "![" not in text:
        return text

    def repl(m):
        alt = (m.group(1) or "").replace('"', "&quot;")
        url = (m.group(2) or "").strip()
        # Whitespace und HTML-Entities im URL bereinigen
        url = url.replace("&amp;", "&")
        return f'<img src="{url}" alt="{alt}"/>'

    return _MD_IMG_RE.sub(repl, text)


def _markdown_links_to_html(text: str) -> str:
    """Konvertiert Markdown-Links [text](url) in <a>-Tags (ohne Image-Variante)."""
    if "[" not in text:
        return text

    def repl(m):
        label = m.group(1) or ""
        url = (m.group(2) or "").strip().replace("&amp;", "&")
        return f'<a href="{url}">{label}</a>'

    return _MD_LINK_RE.sub(repl, text)


# -------- XML Helpers -------- #

def _make_element(tag: str, attrs: dict, text: str | None = None):
    from docx.oxml import OxmlElement
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k) if ":" in k else k, v)
    if text is not None:
        el.text = text
    return el


def _shade_paragraph(paragraph, hex_color: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = _make_element("w:shd", {
        "w:val": "clear",
        "w:color": "auto",
        "w:fill": hex_color,
    })
    pPr.append(shd)


def _add_hyperlink_to_paragraph(paragraph, text: str, url: str,
                                 color: str = "0563C1", underline: bool = True,
                                 font: str = "Arial", size_pt: int = 10) -> None:
    """
    Fuegt einen anklickbaren Hyperlink in einen bestehenden Absatz ein.
    python-docx hat keinen direkten Hyperlink-Support, daher ueber OxmlElement.
    """
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE

    # Relationship in der Document-Part anlegen
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Schriftart
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rPr.append(rFonts)

    # Schriftgroesse (half-points)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size_pt * 2))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(size_pt * 2))
    rPr.append(szCs)

    # Farbe
    col = OxmlElement("w:color")
    col.set(qn("w:val"), color)
    rPr.append(col)

    # Unterstrichen
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _shade_paragraph_cell(cell, hex_color: str) -> None:
    """Setzt Hintergrundfarbe einer Tabellenzelle."""
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    # Existierendes shd entfernen
    for shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)
